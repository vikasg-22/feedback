from flask import Flask, request, jsonify, render_template, session, send_file
from flask_pymongo import PyMongo
from bson.objectid import ObjectId
from flask_cors import CORS
from collections import Counter
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO
import logging
import os



app = Flask(__name__)
CORS(app)  # Enable CORS for cross-origin requests

# MongoDB configuration
import os

app.config["MONGO_URI"] = "mongodb+srv://vikasewp84:vikas123@cluster0.t4mjxqj.mongodb.net/feedback_db?retryWrites=true&w=majority&appName=Cluster0"

mongo = PyMongo(app)


# Define collections
feedback_collection = mongo.db.feedbacks
subjects_collection = mongo.db.subjects
questions_collection = mongo.db.questions


@app.route("/")
def home():
    return render_template("user.html")


@app.route('/admin-login')
def admin_login():
    return render_template("admin.html")

# Rating text-to-number mapping
rating_map = {"Bad": 1, "Average": 2, "Good": 3, "Very Good": 4}
@app.route("/submit-feedback", methods=["POST"])
def submit_feedback():
    data = request.json
    register_number = data.get("register_number")
    semester = data.get("semester")
    branch = data.get("branch")
    subject = data.get("subject")
    form_type = data.get("form")
    year = data.get("year")
    ratings = data.get("ratings")

    if not register_number or len(register_number) < 5:
        return jsonify({"error": "Invalid register number format!"}), 400

    branch_code = register_number[3:5].upper()
    if branch_code != branch.upper():
        return jsonify({"error": f"Register number does not match selected branch ({branch})!"}), 400

    # Check for existing feedback
    existing_feedback = feedback_collection.find_one({
        "register_number": register_number,
        "semester": semester,
        "branch": branch,
        "subject": subject,
        "year": year
    })

    if existing_feedback:
        return jsonify({"error": "Feedback already submitted for this subject and year!"}), 400

        # Check if feedback limit is set and reached
    branch_doc = subjects_collection.find_one({"branch": branch})
    if branch_doc:
        limits = branch_doc.get("feedback_limits", {})
        sem_limit = limits.get(str(semester))  # Ensure semester is a string
        if sem_limit is not None:
            submitted_count = feedback_collection.count_documents({
                "semester": semester,
                "branch": branch,
                "year": year
            })
            if submitted_count >= sem_limit:
                return jsonify({"error": "Feedback limit reached for this semester!"}), 403


    # Save feedback if all checks pass
    feedback_collection.insert_one({
        "register_number": register_number,
        "semester": semester,
        "branch": branch,
        "subject": subject,
        "form": form_type,
        "year": year,
        "ratings": ratings
    })

    return jsonify({"message": "Feedback submitted successfully!"}), 200
    
@app.route('/fetch-feedback', methods=['GET'])
def fetch_feedback():
    try:
        year = request.args.get('year')
        branch = request.args.get('branch')
        semester = request.args.get('semester')
        subject = request.args.get('subject')
        form = request.args.get('form')

        if not all([year, branch, semester, subject, form]):
            return jsonify({"error": "Missing parameters"}), 400

        # Debug: Print received parameters

        # 1. Get questions for this academic year and semester
        questions_doc = questions_collection.find_one({"academic_year": year})
        if not questions_doc:
            print(f"No questions document found for year: {year}")
            return jsonify({"error": "No questions configured for this academic year"}), 404

        semester_key = f"semester_{semester}"
        questions = questions_doc["semesters"].get(semester_key, {}).get(f"{form}_questions", [])
        
        if not questions:
            print(f"No questions found for semester {semester} and form {form}")
            return jsonify({"error": f"No {form} questions found for semester {semester}"}), 404

        print(f"Found questions: {questions}")

        # 2. Get all matching feedback
        feedbacks = list(feedback_collection.find({
            "year": year,
            "branch": branch,
            "semester": semester,
            "subject": subject,
            "form": form
        }))


        if not feedbacks:
            return jsonify({"message": "No feedback available"}), 404

        # 3. Initialize question summary with all questions
        question_summary = {
            q: {
                "Question": q,
                "1 (Bad)": 0,
                "2 (Average)": 0,
                "3 (Good)": 0,
                "4 (Very Good)": 0,
                "Total": 0
            } for q in questions
        }

        # 4. Process each feedback document
        for i, feedback in enumerate(feedbacks):
            print(f"\nProcessing feedback {i+1}:")
            ratings = feedback.get("ratings", [])
            print(f"Raw ratings: {ratings}")

            for q_index, rating in enumerate(ratings):
                if q_index >= len(questions):
                    print(f"Skipping rating index {q_index} - beyond question count")
                    break
                
                question = questions[q_index]
                try:
                    # Convert rating to integer (handle both numeric and text ratings)
                    if isinstance(rating, str):
                        rating = rating_map.get(rating.strip(), 0)
                    rating = int(rating)
                    
                    if 1 <= rating <= 4:
                        label = f"{rating} ({['Bad', 'Average', 'Good', 'Very Good'][rating-1]})"
                        question_summary[question][label] += 1
                        question_summary[question]["Total"] += 1
                        print(f"Question '{question}': Counted rating {rating}")
                    else:
                        print(f"Question '{question}': Invalid rating value {rating}")
                except (ValueError, TypeError) as e:
                    print(f"Question '{question}': Error processing rating {rating}: {str(e)}")
                    continue

        # 5. Calculate averages and prepare response
        feedback_data = []
        for question, data in question_summary.items():
            total_score = (
                data["1 (Bad)"] * 1 +
                data["2 (Average)"] * 2 +
                data["3 (Good)"] * 3 +
                data["4 (Very Good)"] * 4
            )
            total_responses = data["Total"]
            avg = total_score / total_responses if total_responses > 0 else 0
            data["AVG"] = round(avg, 2)
            feedback_data.append(data)
            print(f"Final stats for '{question}': {data}")

        # 6. Get feedback limit
        subject_data = subjects_collection.find_one({'branch': branch})
        feedback_limit = subject_data.get("feedback_limits", {}).get(str(semester), 0) if subject_data else 0

        return jsonify({
            "feedback_data": feedback_data,
            "year": year,
            "branch": branch,
            "semester": semester,
            "subject": subject,
            "form": form,
            "total_submissions": len(feedbacks),
            "feedback_limit": feedback_limit
        })

    except Exception as e:
        print(f"Error in fetch-feedback: {str(e)}")
        return jsonify({"error": "Internal server error", "details": str(e)}), 500

@app.route("/set-feedback-limit", methods=["POST"])
def set_feedback_limit():
    data = request.json
    branch = data.get("branch")
    semester = str(data.get("semester"))  # Ensure it's a string
    limit = data.get("limit")

    if not branch or not semester or not isinstance(limit, int):
        return jsonify({"error": "Missing or invalid data"}), 400

    result = subjects_collection.update_one(
        {"branch": branch},
        {"$set": {f"feedback_limits.{semester}": limit}}
    )

    if result.matched_count == 0:
        return jsonify({"error": "Branch not found"}), 404

    return jsonify({"message": f"Feedback limit for {branch} semester {semester} set to {limit}"}), 200

@app.route("/add-question", methods=["POST"])
def add_question():
    data = request.json
    semester = data.get("semester")
    question = data.get("question")
    qtype = data.get("type")  # "mid" or "end"
    year = data.get("year")

    if not all([semester, question, qtype, year]):
        return jsonify({"message": "Missing data"}), 400

    # Find the existing document for the year
    doc = questions_collection.find_one({"academic_year": year})
    if not doc:
        # Return error if year not found — no creation here
        return jsonify({"message": f"No document found for academic year {year}"}), 404

    # Check if semesters field exists, else initialize it for this semester
    if "semesters" not in doc or semester not in doc["semesters"]:
        # Initialize semester with empty mid_questions and end_questions arrays
        questions_collection.update_one(
            {"academic_year": year},
            {"$set": {f"semesters.{semester}": {"mid_questions": [], "end_questions": []}}}
        )

    # Push new question into the right array under semesters.semester.qtype_questions
    result = questions_collection.update_one(
        {"academic_year": year},
        {"$push": {f"semesters.{semester}.{qtype}_questions": question}}
    )

    if result.modified_count == 0:
        return jsonify({"message": "Failed to add question"}), 500

    return jsonify({"message": "Question added successfully"})

@app.route("/delete-question", methods=["DELETE"])
def delete_question():
    data = request.json
    semester = data.get("semester")  # Should be "semester_X"
    qtype = data.get("type")        # "mid" or "end"
    index = data.get("index")
    year = data.get("year")

    print(f"Received delete request for: {semester}, {qtype}, index {index}, year {year}")  # Debug log

    if not all([semester, qtype, index is not None, year]):
        return jsonify({"message": "Missing required data"}), 400

    # Find the document for this academic year
    doc = questions_collection.find_one({"academic_year": year})
    if not doc:
        return jsonify({"message": "Academic year not found"}), 404

    # Verify the semester exists
    if semester not in doc.get("semesters", {}):
        return jsonify({"message": "Semester not found"}), 404

    # Get the questions array
    questions = doc["semesters"][semester].get(f"{qtype}_questions", [])
    
    # Validate index
    if index < 0 or index >= len(questions):
        return jsonify({"message": "Invalid question index"}), 400

    # Get the exact question text to remove
    question_to_remove = questions[index]
    path = f"semesters.{semester}.{qtype}_questions"

    # Perform the update
    result = questions_collection.update_one(
        {"academic_year": year},
        {"$pull": {path: question_to_remove}}
    )

    if result.modified_count == 0:
        return jsonify({"message": "No changes made - question may not exist"}), 400

    return jsonify({"message": "Question deleted successfully"})

@app.route("/get-questions/<semester>/<form_type>")
def get_questions(semester, form_type):
    year = request.args.get("year")
    if not year:
        return jsonify({"error": "Academic year is required"}), 400

    print(f"Debug - Searching for: year={year}, semester={semester}, form_type={form_type}")  # Debug log

    doc = questions_collection.find_one({"academic_year": year})

    if not doc:
        print("Debug - No document found for year:", year)  # Debug log
        return jsonify({"questions": []})

    try:
        questions = doc["semesters"][semester][f"{form_type}_questions"]
        return jsonify({"questions": questions})
    except KeyError as e:
        print(f"Debug - Key error: {str(e)}")  # Debug log
        return jsonify({"questions": []})

@app.route('/add-subject', methods=['POST'])
def add_subject():
    try:
        data = request.json
        branch = data.get("branch")
        semester = data.get("semester")
        new_subject = data.get("subject")

        if not (branch and semester and new_subject):
            return jsonify({"error": "Branch, Semester, and Subject are required!"}), 400

        # Append the new subject to the array of the specified semester
        subjects_collection.update_one(
            {"branch": branch},
            {"$push": {f"semesters.{semester}": new_subject}},
            upsert=True
        )

        return jsonify({"message": "Subject added successfully!"}), 201
    except Exception as e:
        print(f"Error: {e}")  # Debugging line
        return jsonify({"error": str(e)}), 500



@app.route('/get-subjects/<branch>/<semester>', methods=['GET'])
def get_subjects(branch, semester):
    subjects_data = subjects_collection.find_one({"branch": branch})
    
    if not subjects_data or semester not in subjects_data["semesters"]:
        return jsonify({"error": "No subjects found for this semester."}), 404

    subjects_list = subjects_data["semesters"][semester]
    return jsonify(subjects_list), 200

@app.route('/get-all-subjects', methods=['GET'])
def get_all_subjects():
    # Fetch all subjects from the database
    all_subjects = subjects_collection.find()
    subjects_data = []
    
    for subject_doc in all_subjects:
        branch = subject_doc.get("branch", "Unknown Branch")
        semesters = subject_doc.get("semesters", {})
        
        formatted_subjects = []
        for semester, subjects in semesters.items():
            formatted_subjects.append({
                "semester": semester,
                "subjects": subjects
            })
        
        subjects_data.append({
            "branch": branch,
            "semesters": formatted_subjects
        })
    
    return jsonify(subjects_data), 200
    
@app.route('/get-branches', methods=['GET'])
def get_branches():
    branches = subjects_collection.distinct("branch")
    return jsonify(branches), 200

@app.route('/delete-subject', methods=['DELETE'])
def delete_subject():
    data = request.json
    branch = data.get('branch')
    semester = str(data.get('semester'))  # Ensure semester is a string
    subject = data.get('subject')

    if not branch or not semester or not subject:
        return jsonify({"success": False, "message": "Missing required fields!"}), 400

    # Find the document for the specified branch
    subject_doc = subjects_collection.find_one({"branch": branch})
    
    if not subject_doc:
        return jsonify({"success": False, "message": "Branch not found!"}), 404

    # Check if the semester exists in the document
    if semester not in subject_doc.get("semesters", {}):
        return jsonify({"success": False, "message": "Semester not found!"}), 404

    # Check if the subject exists in the semester list
    if subject not in subject_doc["semesters"][semester]:
        return jsonify({"success": False, "message": "Subject not found!"}), 404

    # Remove the subject from the semester list
    subjects_collection.update_one(
        {"branch": branch},
        {"$pull": {f"semesters.{semester}": subject}}
    )

    return jsonify({"success": True, "message": "Subject deleted successfully!"}), 200

@app.route('/get-all-questions', methods=['GET'])
def get_all_questions():
    # Fetch all questions from the database
    question_doc = questions_collection.find_one({"_id": "questions_collection"})
    if not question_doc or "questions" not in question_doc:
        return jsonify({"error": "No questions found!"}), 404

    return jsonify(question_doc["questions"]), 200

@app.route('/download-feedback', methods=['GET'])
def download_feedback():
    branch = request.args.get("branch")
    semester = request.args.get("semester")
    subject = request.args.get("subject")
    form = request.args.get("form")  # 'mid' or 'end'
    year = request.args.get("year")

    if not all([branch, semester, subject, form, year]):
        return jsonify({"error": "Missing parameters"}), 400

    feedback_data = list(feedback_collection.find(
        {"branch": branch, "semester": semester, "subject": subject, "form": form, "year": year},
        {"_id": 0, "ratings": 1}
    ))

    if not feedback_data:
        return jsonify({"message": "No feedback available"}), 404

    academic_year = year.split("-")[0]
    next_year = str(int(academic_year) + 1)
    academic_year_str = f"{academic_year}-{next_year}"

    question_doc = questions_collection.find_one({"academic_year": academic_year_str})
    if not question_doc:
        return jsonify({"error": f"No questions found for academic year {academic_year_str}!"}), 404

    semester_key = f"semester_{semester}"
    if semester_key not in question_doc.get("semesters", {}):
        return jsonify({"error": f"No questions found for semester {semester}!"}), 404

    questions_key = f"{form}_questions"
    questions = question_doc["semesters"][semester_key].get(questions_key, [])

    if not questions:
        return jsonify({"error": f"No {form} questions found for semester {semester}!"}), 404

    print("Questions fetched:", questions[:2])  # debug first 2 questions

  

    try:
        file_stream = create_feedback_template(subject, semester, branch, year, feedback_data, questions, form)
        if file_stream is None:
            raise ValueError("Failed to generate document")
        
        return send_file(
            file_stream,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name=f"feedback_{branch}{semester}{subject}_{form}_{year}.docx"
        )
    except Exception as e:

        return jsonify({"error": f"Failed to generate document: {str(e)}"}), 500

def create_feedback_template(subject, semester, branch, year, feedback_data, questions, formtype):
    doc = Document()
    
    # Add basic document properties
    doc.core_properties.title = f"Feedback Report - {subject}"
    doc.core_properties.author = "Feedback System"
    
    # 1. HEADER SECTION
    header = doc.sections[0].header
    hdr_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hdr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hdr_run = hdr_para.add_run("STUDENT FEEDBACK FORM\n")
    hdr_run.bold = True
    hdr_run.font.size = Pt(14)
    
    # 2. INSTITUTION DETAILS
    doc.add_paragraph("Institution Name: East West Polytechnic", style='Heading2')
    doc.add_paragraph("Institution Code: 499", style='Heading2')
    doc.add_paragraph()  # Empty paragraph for spacing
    
    # 3. METADATA TABLE
    table_details = doc.add_table(rows=1, cols=2)
    table_details.style = 'Table Grid'
    
    # Left cell
    left_cell = table_details.rows[0].cells[0]
    left_cell.text = f"Branch: {branch}\nSemester: {semester}\nSubject: {subject}"
    
    # Right cell
    right_cell = table_details.rows[0].cells[1]
    right_cell.text = f"Form Type: {formtype}\nYear: {year}\nTeacher's Name: ________"
    right_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_paragraph()  # Spacing
    
    # 4. FEEDBACK TABLE
    if not questions or not feedback_data:
        doc.add_paragraph("No feedback data available", style='Intense Quote')
    else:
        table = doc.add_table(rows=1, cols=7)
        table.style = 'Table Grid'
        
        # Set column widths
        widths = [0.5, 3.5, 0.8, 0.8, 0.8, 0.8, 0.8]  # in inches
        for i, width in enumerate(widths):
            table.columns[i].width = Inches(width)
        
        # Header row
        hdr_cells = table.rows[0].cells
        headers = ["S.No", "Question", "1 (Bad)", "2 (Average)", "3 (Good)", "4 (Very Good)", "Average"]
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            hdr_cells[i].paragraphs[0].runs[0].bold = True
            hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Process each question
        total_avg_sum = 0
        valid_questions = 0
        
        for i, question in enumerate(questions):
            # Extract question text
            if isinstance(question, dict):
                q_text = question.get('text', f'Question {i+1}')
            else:
                q_text = str(question)
            
            # Initialize rating counters
            rating_counts = {1: 0, 2: 0, 3: 0, 4: 0}
            total_responses = 0
            total_rating = 0
            
            # Process each feedback response
            for feedback in feedback_data:
                if i < len(feedback.get('ratings', [])):
                    rating = feedback['ratings'][i]
                    
                    # Convert rating to numeric if needed
                    if isinstance(rating, str):
                        rating = {'Bad': 1, 'Average': 2, 'Good': 3, 'Very Good': 4}.get(rating, 0)
                    elif not isinstance(rating, (int, float)):
                        rating = 0
                    
                    if 1 <= rating <= 4:
                        rating_counts[int(rating)] += 1
                        total_responses += 1
                        total_rating += rating
            
            # Calculate average for this question
            avg = total_rating / total_responses if total_responses > 0 else 0
            total_avg_sum += avg
            valid_questions += 1 if total_responses > 0 else 0
            
            # Add row to table
            row_cells = table.add_row().cells
            row_cells[0].text = str(i+1)
            row_cells[1].text = q_text
            row_cells[2].text = str(rating_counts[1])
            row_cells[3].text = str(rating_counts[2])
            row_cells[4].text = str(rating_counts[3])
            row_cells[5].text = str(rating_counts[4])
            row_cells[6].text = f"{avg:.2f}"
            
            # Center align numeric cells
            for j in range(2, 7):
                row_cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add total average
        if valid_questions > 0:
            total_avg = total_avg_sum / valid_questions
            doc.add_paragraph()
            para = doc.add_paragraph()
            para.add_run(f"Overall Average Rating: {total_avg:.2f}").bold = True
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Save to memory
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream  
if __name__ == "__main__":
    app.run(debug=True)



